# Word 模板填充引擎
import os
from docxtpl import DocxTemplate
from typing import Dict, Any
from src.utils.logger_handler import get_logger

logger = get_logger(__name__)


class WordEngine:
    """Word 模板填充引擎"""
    
    def __init__(self, templates_dir: str) -> None:
        """初始化 Word 引擎"""
        # 如果是相对路径，转换为绝对路径（相对于项目根目录）
        if not os.path.isabs(templates_dir):
            # 获取项目根目录
            project_root = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
            templates_dir = os.path.join(project_root, templates_dir)
        
        self.templates_dir = templates_dir
        # 确保模板目录存在
        os.makedirs(self.templates_dir, exist_ok=True)
        logger.info(f"Word 引擎初始化完成，模板目录: {self.templates_dir}")
    
    def render(self, template_name: str, context: Dict[str, Any]) -> str:
        """渲染 Word 模板"""
        try:
            # 构建模板路径
            template_path = os.path.join(self.templates_dir, template_name)
            logger.debug(f"渲染模板: {template_path}")
            
            # 检查模板是否存在
            if not os.path.exists(template_path):
                # 使用默认模板
                template_path = os.path.join(self.templates_dir, 'default_template.docx')
                logger.debug(f"使用默认模板: {template_path}")
                
                if not os.path.exists(template_path):
                    # 创建默认模板
                    logger.info(f"创建默认模板: {template_path}")
                    self._create_default_template(template_path)
            
            # 加载模板
            doc = DocxTemplate(template_path)
            logger.debug("模板加载成功")
            
            # 准备上下文数据
            render_context = self._prepare_context(context)
            
            # 渲染模板
            doc.render(render_context)
            logger.debug("模板渲染成功")
            
            # 生成输出路径
            output_dir = os.path.join(os.path.dirname(self.templates_dir), 'output')
            os.makedirs(output_dir, exist_ok=True)
            
            output_filename = f"generated_{template_name}"
            output_path = os.path.join(output_dir, output_filename)
            
            # 保存文档
            doc.save(output_path)
            logger.info(f"文档保存成功: {output_path}")
            
            return output_path
        except Exception as e:
            logger.error(f"Word 渲染错误: {str(e)}")
            return ""
    
    def _prepare_context(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """准备渲染上下文"""
        render_context = {
            'title': context.get('title', '文档'),
            'sections': context.get('sections', []),
            'metadata': context.get('metadata', {})
        }
        
        # 处理章节内容
        for section in render_context['sections']:
            # 处理子章节
            if 'subsections' in section:
                for subsection in section['subsections']:
                    subsection['content'] = self._format_content(subsection.get('content', ''))
            section['content'] = self._format_content(section.get('content', ''))
        
        return render_context
    
    def _format_content(self, content: str) -> str:
        """格式化内容"""
        # 简单的内容格式化
        return content
    
    def _create_default_template(self, template_path: str) -> None:
        """创建默认模板"""
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('{{title}}', 0)
            doc.add_paragraph('标题: {{title}}')
            doc.add_paragraph('文档类型: {{metadata.doc_type}}')
            doc.add_paragraph('\n章节内容:')
            doc.add_paragraph('{% for section in sections %}')
            doc.add_paragraph('{{section.level}}. {{section.title}}')
            doc.add_paragraph('{{section.content}}')
            doc.add_paragraph('{% endfor %}')
            doc.save(template_path)
            logger.info(f"默认模板创建成功: {template_path}")
        except Exception as e:
            logger.error(f"创建默认模板失败: {str(e)}")
            # 创建一个简单的文本文件作为备选
            with open(template_path, 'w', encoding='utf-8') as f:
                f.write('{{title}}\n\n{{content}}')
            logger.warning(f"已创建简单文本模板作为备选: {template_path}")
