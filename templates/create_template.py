"""
创建默认的 Word 模板文件
"""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


def create_default_template():
    """创建默认的 Word 模板"""
    # 获取项目根目录
    current_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    templates_dir = os.path.join(current_dir, 'templates')
    template_path = os.path.join(templates_dir, 'default_template.docx')

    # 确保目录存在
    os.makedirs(templates_dir, exist_ok=True)

    # 创建文档
    doc = Document()

    # 添加标题
    title = doc.add_heading('文档标题', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # 添加元信息段落
    meta = doc.add_paragraph()
    meta.add_run('文档类型: ').bold = True
    meta.add_run('分析报告')
    meta.add_run('\n生成时间: ').bold = True
    meta.add_run('2024年')

    # 添加分隔线
    doc.add_paragraph('=' * 60)

    # 添加章节标题
    doc.add_heading('章节内容', level=1)

    # 添加示例章节
    section = doc.add_paragraph()
    section.add_run('1. 章节标题').bold = True
    doc.add_paragraph('这里是章节的详细内容...')

    # 保存文档
    doc.save(template_path)
    print(f"模板创建成功: {template_path}")
    return template_path


if __name__ == "__main__":
    create_default_template()
